"""
Generate WeChat Official Article in DOCX format for PlotBase MCP with embedded images & contribution notice.
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Set shading color for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_code_block(doc, code_text):
    """Add a beautifully styled code block with background shading and left border."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F4F6F8")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="0D9488"/>'
        f'<w:top w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def add_callout_box(doc, title, text, border_color="0284C7", bg_color="F0F9FF", title_color="0369A1"):
    """Add a highlighted callout box for key takeaways or announcements."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'<w:top w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    
    run_title = p.add_run(f"💡 {title}\n")
    run_title.font.name = 'Arial'
    run_title.font.bold = True
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = RGBColor.from_string(title_color)

    run_text = p.add_run(text)
    run_text.font.name = '宋体'
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(6)

def add_image_with_caption(doc, image_path, caption, width=Inches(5.8)):
    """Embed an image with a centered caption."""
    if not os.path.exists(image_path):
        print(f"Warning: image path {image_path} not found")
        return

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(4)
    
    run_img = p_img.add_run()
    run_img.add_picture(image_path, width=width)

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    
    r_cap = p_cap.add_run(f"▲ {caption}")
    r_cap.font.name = '宋体'
    r_cap.font.size = Pt(9.5)
    r_cap.font.italic = True
    r_cap.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

def build_document():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = '宋体'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("告别审美灾难！PlotBase-MCP 让 AI Agent 秒变顶刊科研可视化大牛 🎨")
    r_title.font.name = '微软雅黑'
    r_title.font.bold = True
    r_title.font.size = Pt(20)
    r_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    r_sub = p_sub.add_run("基于 Model Context Protocol 的科研图表制作档案与 Agent 绘图上下文管理系统")
    r_sub.font.name = '微软雅黑'
    r_sub.font.italic = True
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    # Meta banner
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(18)
    r_meta = p_meta.add_run("作者：Antigravity AI Team | GitHub 开源项目 | 阅读时间：约 8 分钟")
    r_meta.font.name = 'Arial'
    r_meta.font.size = Pt(9.5)
    r_meta.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    # Notice callout at top
    add_callout_box(
        doc,
        "补充说明 & 示例征集 (Call for Examples)",
        "📢 PlotBase 的图表制作档案库目前仍处于【持续完善与扩充中】！\n"
        "我们极其欢迎广大科研人员、数据分析师与开源爱好者共同建设。如果你手头有高质量的 R/Python 绘图代码、顶刊配色模版或特定图表的 Agent 微调经验，非常欢迎提交示例与建议！",
        border_color="D97706", bg_color="FFFBEB", title_color="B45309" # Warm Amber
    )

    # Separator
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep.add_run("─────────── ✦ ───────────").font.color.rgb = RGBColor(0xCB, 0xD5, 0xE0)
    p_sep.paragraph_format.space_after = Pt(18)

    # Section 1: 痛点与背景
    h1 = doc.add_paragraph()
    r_h1 = h1.add_run("一、为什么普通的 AI Prompt 画不出合格的科研图表？")
    r_h1.font.name = '微软雅黑'
    r_h1.font.bold = True
    r_h1.font.size = Pt(15)
    r_h1.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(8)

    p1 = doc.add_paragraph(
        "在日常科研工作中，利用 LLM（大语言模型）生成 R 或 Python 绘图代码已成为许多研究者的习惯。然而，大家普遍遇到以下三大痛点："
    )
    p1.paragraph_format.space_after = Pt(6)

    bullets = [
        ("审美灾难与排版混乱：", "大模型默认生成的图表配色往往极其刺眼，字号比例失调，标签文字互相重叠，难以达到 Nature/Cell 等顶刊发表标准。"),
        ("缺乏复杂图表上下文：", "对于单细胞拟时序热图、多组山脊图、云雨图、曼哈顿图等高阶可视化，模型往往给错依赖包或遗漏数据标准化步骤。"),
        ("列名映射与调试噩梦：", "模型生成的代码死板硬套示例字段，研究者拿到代码后需要手动逐行修改变量名、调整 Scale 和修改色盘，极其耗费时间。")
    ]
    for b_title, b_desc in bullets:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(4)
        r_bt = p_b.add_run(b_title)
        r_bt.font.bold = True
        r_bt.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)
        p_b.add_run(b_desc)

    add_callout_box(
        doc,
        "核心突破：PlotBase-MCP 的解法",
        "通过 Model Context Protocol (MCP)，将图表从单纯的『一段代码』升级为『5维制作档案 (Figure Archive)』，包含元数据、完整代码、示例数据、顶刊视觉规则及专门针对 Agent 的修改指导，让 AI Agent 在绘图前拥有完全的领域知识上下文！"
    )

    # Section 2: 制作档案 5 维结构
    h2 = doc.add_paragraph()
    r_h2 = h2.add_run("二、解密 PlotBase-MCP 的 5 维图表制作档案与架构")
    r_h2.font.name = '微软雅黑'
    r_h2.font.bold = True
    r_h2.font.size = Pt(15)
    r_h2.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(8)

    doc.add_paragraph("PlotBase MCP 数据库将科研图表抽象为包含 5 个核心维度的结构化档案：")

    add_image_with_caption(
        doc,
        "/tmp/plotbase_mcp/images/fig1_architecture.png",
        "图 1：PlotBase-MCP 架构设计与 5 维制作档案交互流程图",
        width=Inches(6.0)
    )

    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    headers = ["维度名称", "包含内容", "Agent 获得的收益"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0D9488")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    data_rows = [
        ("📌 1. 元数据 (Metadata)", "标题、分类大类、图表类型、R/Python依赖包、标签", "精确检索与环境依赖智能提示"),
        ("💻 2. 示例代码 (Code)", "经过严格验证、可直接运行的 R (ggplot2) / Python 源码", "获得高质量无语法错误的底层绘图逻辑"),
        ("📊 3. 示例数据 (Data)", "CSV / JSON 样例数据格式、列名与示例数值", "明确数据输入格式，方便匹配研究数据"),
        ("🎨 4. 视觉规则 (Rules)", "顶刊色盘 (Viridis/Set2)、字体、主题、比例", "强制约束配色与排版，杜绝审美灾难"),
        ("🤖 5. Agent 上下文", "专用修改指导、变量映射规则、数据量阈值建议", "指导 Agent 自动完成列名替换与动态调参")
    ]

    col_widths = [Inches(1.8), Inches(2.7), Inches(2.0)]
    for row_idx, row_data in enumerate(data_rows, start=1):
        row_cells = table.rows[row_idx].cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            row_cells[col_idx].width = col_widths[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 3: 实战案例与图表效果展示
    h3 = doc.add_paragraph()
    r_h3 = h3.add_run("三、顶刊级科研图表渲染案例与代码展示")
    r_h3.font.name = '微软雅黑'
    r_h3.font.bold = True
    r_h3.font.size = Pt(15)
    r_h3.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    h3.paragraph_format.space_before = Pt(16)
    h3.paragraph_format.space_after = Pt(8)

    # Case 1: 山脊图
    p_c1 = doc.add_paragraph()
    r_c1_t = p_c1.add_run("案例 1：多组连续数据分布山脊图 (ggplot2 + ggridges)\n")
    r_c1_t.font.bold = True
    r_c1_t.font.size = Pt(12)
    r_c1_t.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)
    p_c1.add_run("当研究包含多个平行处理组（>3组）且需展现连续分布趋势时，山脊图（Ridge Plot）是表现力极佳的选择：")

    add_image_with_caption(
        doc,
        "/tmp/plotbase_mcp/images/fig2_ridge_plot.png",
        "图 2：PlotBase-MCP 预置案例 —— 基于 viridis 配色与 ggridges 的多组连续分布山脊图效果",
        width=Inches(5.5)
    )

    code_ridge = """library(ggplot2)
library(ggridges)
library(viridis)

# 加载研究数据
df <- read.csv("sample_density_data.csv")

# 绘制多组连续分布山脊图
ggplot(df, aes(x = value, y = group, fill = ..x..)) +
  geom_density_ridges_gradient(scale = 3, rel_min_height = 0.01) +
  scale_fill_viridis(name = "Value", option = "C") +
  labs(title = '多组连续数据分布山脊图', subtitle = '基于 ggridges 与 viridis 色盘',
       x = '测定值 (Value)', y = '分组 (Group)') +
  theme_ridges(font_size = 13, grid = TRUE) +
  theme(legend.position = "right", panel.spacing = unit(0.1, "lines"))"""
    add_code_block(doc, code_ridge)

    # Case 2: 火山图
    p_c2 = doc.add_paragraph()
    r_c2_t = p_c2.add_run("案例 2：高发表级转录组差异表达火山图 (EnhancedVolcano)\n")
    r_c2_t.font.bold = True
    r_c2_t.font.size = Pt(12)
    r_c2_t.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)
    p_c2.add_run("转录组/蛋白组差异表达分析的标准可视化方案，清晰标记显著上调、下调基因及关键基因 Tag：")

    add_image_with_caption(
        doc,
        "/tmp/plotbase_mcp/images/fig3_volcano_plot.png",
        "图 3：PlotBase-MCP 预置案例 —— 自动标注上/下调差异基因的发表级火山图渲染效果",
        width=Inches(5.2)
    )

    code_volcano = """library(EnhancedVolcano)

EnhancedVolcano(res,
  lab = rownames(res),
  x = 'log2FoldChange',
  y = 'pvalue',
  pCutoff = 10e-6,
  FCcutoff = 1.5,
  pointSize = 3.0,
  labSize = 4.0,
  col=c('black', 'black', 'blue', 'red'),
  colAlpha = 0.8,
  legendPosition = 'right',
  title = '差异基因火山图 (Volcano Plot)',
  subtitle = 'Significant Up/Down Regulated Genes'
)"""
    add_code_block(doc, code_volcano)

    # Case 3: 时间序列图
    p_c3 = doc.add_paragraph()
    r_c3_t = p_c3.add_run("案例 3：多指标动态时间序列走势图 (带 95% 置信区间阴影)\n")
    r_c3_t.font.bold = True
    r_c3_t.font.size = Pt(12)
    r_c3_t.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)
    p_c3.add_run("适用于动物体重、基因表达随时间变化或临床随访等连续观测数据的折线趋势呈现：")

    add_image_with_caption(
        doc,
        "/tmp/plotbase_mcp/images/fig4_line_ci_plot.png",
        "图 4：PlotBase-MCP 预置案例 —— Python (Seaborn) 带 95% 置信区间阴影的多组时间序列折线图",
        width=Inches(5.5)
    )

    code_line = """import matplotlib.pyplot as plt
import seaborn as sns

sns.set_theme(style="whitegrid")
plt.figure(figsize=(8, 5), dpi=300)

ax = sns.lineplot(
    data=df, x="date", y="value", hue="category",
    style="category", markers=True, dashes=False,
    errorbar=("ci", 95), linewidth=2.5
)
plt.title("多指标时间序列动态走势", fontsize=14, fontweight='bold', pad=15)
plt.xlabel("时间 (Time)", fontsize=12)
plt.ylabel("测量指标 (Metric)", fontsize=12)
plt.legend(title="类别", frameon=True)
plt.tight_layout()"""
    add_code_block(doc, code_line)

    # Section 4: 接入指南
    h4 = doc.add_paragraph()
    r_h4 = h4.add_run("四、如何快速接入 PlotBase-MCP 服务？")
    r_h4.font.name = '微软雅黑'
    r_h4.font.bold = True
    r_h4.font.size = Pt(15)
    r_h4.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    h4.paragraph_format.space_before = Pt(16)
    h4.paragraph_format.space_after = Pt(8)

    doc.add_paragraph("只需两步即可在 Claude Desktop、Cursor 或 Windsurf 中启用科研图表档案库：")

    p_step1 = doc.add_paragraph(style='List Number')
    p_step1.add_run("克隆 GitHub 仓库并安装依赖：").font.bold = True
    add_code_block(doc, "git clone https://github.com/Avatar-stack/plotbase-mcp.git\ncd plotbase-mcp\npip install -e .")

    p_step2 = doc.add_paragraph(style='List Number')
    p_step2.add_run("配置客户端 MCP Settings (例如 claude_desktop_config.json)：").font.bold = True
    code_mcp_config = """{
  "mcpServers": {
    "plotbase-figurememory": {
      "command": "python3",
      "args": [
        "/path/to/plotbase-mcp/server.py"
      ]
    }
  }
}"""
    add_code_block(doc, code_mcp_config)

    # Section 5: 总结与示例征集
    h5 = doc.add_paragraph()
    r_h5 = h5.add_run("五、总结与示例征集致谢")
    r_h5.font.name = '微软雅黑'
    r_h5.font.bold = True
    r_h5.font.size = Pt(15)
    r_h5.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    h5.paragraph_format.space_before = Pt(16)
    h5.paragraph_format.space_after = Pt(8)

    p_end = doc.add_paragraph(
        "PlotBase-MCP 的目标是弥合大语言模型与高水准科研可视化之间的审美与技术鸿沟。让研究者无需在琐碎的代码语法和配色参数上浪费大量精力，只需向 Agent 表达研究需求，即可获得发表级别的优雅图表。"
    )
    p_end.paragraph_format.space_after = Pt(12)

    # Final callout for community contribution
    add_callout_box(
        doc,
        "欢迎提交示例与共同建设",
        "目前 PlotBase 记忆库的案例还在持续完善与扩充中！欢迎广大科研作者与开发者提供高质量绘图示例：\n"
        "• 提交 PR 添加全新的 R/Python 制作档案；\n"
        "• 提交 Issue 推荐优雅的配色规则与 Agent 调优经验。\n\n"
        "🔗 GitHub 开源项目地址：https://github.com/Avatar-stack/plotbase-mcp\n"
        "开源许可：MIT License",
        border_color="0D9488", bg_color="F0FDF4", title_color="0F766E"
    )

    output_path = "/tmp/plotbase_mcp/PlotBase_MCP_微信公众号文章.docx"
    doc.save(output_path)
    print(f"Successfully updated DOCX document with contribution notice at {output_path}")

if __name__ == "__main__":
    build_document()
